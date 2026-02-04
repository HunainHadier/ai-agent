import csv
import os
from datetime import date, datetime
from typing import Iterable

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader
import xlrd

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings

from app.config import get_openai_api_key
from app.storage import agent_vector_dir


SUPPORTED_EXTENSIONS = {
    ".txt",
    ".md",
    ".json",
    ".yaml",
    ".yml",
    ".py",
    ".csv",
    ".pdf",
    ".docx",
    ".xlsx",
    ".xls",
}


def _stringify_cell(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, (datetime, date)):
        return value.isoformat()
    if isinstance(value, float):
        if value.is_integer():
            return str(int(value))
        return str(value)
    text = str(value).strip()
    return text


def _read_pdf_file(path: str) -> str | None:
    try:
        reader = PdfReader(path)
    except Exception:
        return None
    parts: list[str] = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if text.strip():
            parts.append(text)
    content = "\n".join(parts).strip()
    return content or None


def _read_docx_file(path: str) -> str | None:
    try:
        doc = Document(path)
    except Exception:
        return None
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    content = "\n".join(parts).strip()
    return content or None


def _read_xlsx_file(path: str) -> str | None:
    try:
        workbook = load_workbook(path, read_only=True, data_only=True)
    except Exception:
        return None
    parts: list[str] = []
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows(values_only=True):
            cells = [_stringify_cell(cell) for cell in row]
            cells = [cell for cell in cells if cell != ""]
            if cells:
                parts.append(" | ".join(cells))
    content = "\n".join(parts).strip()
    return content or None


def _read_xls_file(path: str) -> str | None:
    try:
        workbook = xlrd.open_workbook(path)
    except Exception:
        return None
    parts: list[str] = []
    for sheet in workbook.sheets():
        for rowx in range(sheet.nrows):
            row_cells: list[str] = []
            for colx in range(sheet.ncols):
                cell = sheet.cell(rowx, colx)
                value = cell.value
                if cell.ctype == xlrd.XL_CELL_DATE:
                    try:
                        value = xlrd.xldate_as_datetime(value, workbook.datemode)
                    except Exception:
                        value = cell.value
                cell_text = _stringify_cell(value)
                if cell_text != "":
                    row_cells.append(cell_text)
            if row_cells:
                parts.append(" | ".join(row_cells))
    content = "\n".join(parts).strip()
    return content or None


def _read_csv_file(path: str) -> str | None:
    try:
        rows: list[str] = []
        with open(path, newline="", encoding="utf-8", errors="ignore") as handle:
            reader = csv.reader(handle)
            for row in reader:
                cells = [str(cell).strip() for cell in row if cell is not None]
                cells = [cell for cell in cells if cell != ""]
                if cells:
                    rows.append(" | ".join(cells))
        content = "\n".join(rows).strip()
        return content or None
    except Exception:
        return None


def _read_text_file(path: str) -> str | None:
    _, ext = os.path.splitext(path)
    ext = ext.lower()
    if ext == ".pdf":
        return _read_pdf_file(path)
    if ext == ".docx":
        return _read_docx_file(path)
    if ext == ".xlsx":
        return _read_xlsx_file(path)
    if ext == ".xls":
        return _read_xls_file(path)
    if ext == ".csv":
        return _read_csv_file(path)
    if ext not in SUPPORTED_EXTENSIONS:
        return None
    with open(path, "r", encoding="utf-8", errors="ignore") as handle:
        return handle.read()


def collect_texts(files: Iterable[str]) -> list[str]:
    texts = []
    for path in files:
        if not os.path.isfile(path):
            continue
        content = _read_text_file(path)
        if content:
            texts.append(content)
    return texts


def build_vectorstore(name: str, texts: Iterable[str], device_id: str | int | None = None) -> str:
    api_key = get_openai_api_key()
    embeddings = OpenAIEmbeddings(api_key=api_key)
    splitter = RecursiveCharacterTextSplitter(chunk_size=900, chunk_overlap=150)
    docs = splitter.create_documents(list(texts))
    vector_dir = agent_vector_dir(device_id, name)
    if docs:
        store = FAISS.from_documents(docs, embeddings)
        store.save_local(vector_dir)
    return vector_dir


def load_vectorstore(name: str, device_id: str | int | None = None) -> FAISS:
    api_key = get_openai_api_key()
    embeddings = OpenAIEmbeddings(api_key=api_key)
    vector_dir = agent_vector_dir(device_id, name)
    return FAISS.load_local(vector_dir, embeddings, allow_dangerous_deserialization=True)
